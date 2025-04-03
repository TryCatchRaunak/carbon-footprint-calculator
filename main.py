import chromadb

chroma_client = chromadb.HttpClient(host="0.0.0.0", port=8000)  # Forces in-memory mode


import streamlit as st
import json
import io
import os
import matplotlib.pyplot as plt
from dotenv import load_dotenv
from docx import Document
from crewai import Agent, Task, Crew, Process

# Load environment variables
load_dotenv()

# Streamlit UI Configuration
st.set_page_config(page_title="Carbon Footprint Calculator", page_icon="🌍", layout="wide")
st.title("🌍 Carbon Footprint Calculator")

# Description
st.markdown("""
### What is a Carbon Footprint?
A carbon footprint is the total amount of greenhouse gases (including carbon dioxide and methane) that are generated by our actions. It is measured in kilograms or tons of CO₂ equivalent.

### Why Calculate Your Carbon Footprint?
Understanding your carbon footprint helps you identify areas where you can reduce your environmental impact. By making small changes in your daily habits, you can contribute to a more sustainable future.

### What Happens After Calculation?
Once your carbon footprint is calculated, you will receive a sustainability score and personalized recommendations to reduce your emissions. These recommendations are tailored to your lifestyle and can help you make eco-friendly choices.
""")

st.markdown("### Reduce Your Environmental Impact 🌱")

# Function to calculate environmental impact score
def calculate_environment_score(total_co2):
    if total_co2 < 50:
        return 90  # Excellent sustainability
    elif total_co2 < 100:
        return 75  # Good sustainability
    elif total_co2 < 200:
        return 50  # Moderate sustainability
    elif total_co2 < 500:
        return 30  # Poor sustainability
    else:
        return max(5, 100 - (total_co2 / 10))  # Very poor sustainability

# Function to create and download a Word document
def create_docx_report(user_data, recommendations, score):
    doc = Document()
    doc.add_heading("🌍 Personalized Carbon Footprint Report", 0)

    doc.add_heading("📊 User Inputs", level=1)
    for category, details in user_data.items():
        doc.add_heading(category.capitalize(), level=2)
        if isinstance(details, dict):
            for key, value in details.items():
                doc.add_paragraph(f"{key.replace('_', ' ').capitalize()}: {value}")
        else:
            doc.add_paragraph(f"{category.replace('_', ' ').capitalize()}: {details}")

    doc.add_heading("💨 Carbon Footprint Breakdown", level=1)
    doc.add_paragraph(f"Estimated Monthly CO₂ Emissions: {user_data.get('total_co2', 'N/A')} kg")

    doc.add_heading("🔄 Recommendations", level=1)
    for section, points in recommendations.items():
        doc.add_heading(section, level=2)
        for point in points:
            doc.add_paragraph(f"• {point}")

    doc.add_heading("♻️ Environmental Impact Score", level=1)
    doc.add_paragraph(f"Your sustainability score: {score}/100")

    # Ensure buffer is properly written and read
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Reset buffer position before returning
    return buffer


# User Input Sections

st.header("🚗 Travel")
flights = st.number_input("✈️ Flights per year", min_value=0, value=2)
flight_distance = st.number_input("🛫 Total flight distance (km)", min_value=0, value=5000)
car_fuel = st.number_input("⛽ Fuel used per month (liters)", min_value=0, value=50)
public_transport_km = st.number_input("🚌 Distance traveled by public transport (km/month)", min_value=0, value=300)
bike_km = st.number_input("🚴 Distance traveled by bike (km/month)", min_value=0, value=50)
walking_km = st.number_input("🚶 Distance walked (km/month)", min_value=0, value=20)

st.header("⚡ Energy Usage")
electricity_kwh = st.number_input("🔋 Monthly electricity usage (kWh)", min_value=0, value=200)
gas_therms = st.number_input("🔥 Monthly gas usage (therms)", min_value=0, value=30)
solar_percentage = st.slider("☀️ % Energy from renewable sources", 0, 100, 20)
heating_kwh = st.number_input("🌡️ Monthly heating energy usage (kWh)", min_value=0, value=100)

st.header("💧 Water Usage")
water_liters = st.number_input("🚰 Monthly water consumption (liters)", min_value=100, value=5000)
hot_water_percentage = st.slider("🔥 % of water heated", 0, 100, 50)

st.header("🍽️ Diet & Consumption")
diet_type = st.selectbox("🥗 Diet type", ["Omnivore", "Vegetarian", "Vegan"])
food_waste_kg = st.number_input("🍲 Monthly food waste (kg)", min_value=0, value=5)
plastic_waste_kg = st.number_input("🛍️ Monthly plastic waste (kg)", min_value=0, value=2)
local_food_percentage = st.slider("🌾 % of food sourced locally", 0, 100, 30)

st.header("🏠 Household & Waste Management")
recycle = st.radio("♻️ Do you recycle?", ["Yes", "No"])
household_size = st.number_input("👨‍👩‍👧 Household size", min_value=1, value=2)
compost = st.radio("🌱 Do you compost organic waste?", ["Yes", "No"])

if st.button("Calculate Footprint"):
    # More realistic carbon footprint calculations
    travel_co2 = (flights * flight_distance * 0.09) + (car_fuel * 2.3) + (public_transport_km * 0.05)
    energy_co2 = ((electricity_kwh * 0.2) + (gas_therms * 0.005) + (heating_kwh * 0.2)) * ((100 - solar_percentage) / 100)
    water_co2 = water_liters * 0.002 * (hot_water_percentage / 100)
    diet_co2 = 30 if diet_type == "Omnivore" else (20 if diet_type == "Vegetarian" else 10)
    waste_co2 = (food_waste_kg * 1.5) + (plastic_waste_kg * 2) + (5 if recycle == "Yes" else 10) + (0 if compost == "Yes" else 5)

    total_co2 = round(travel_co2 + energy_co2 + water_co2 + diet_co2 + waste_co2, 2)
    score = calculate_environment_score(total_co2)

    # CrewAI Agents
    co2_analyst = Agent(role="CO₂ Analyst", goal="Calculate carbon footprint", backstory="Expert in carbon footprint calculations.", verbose=True)
    impact_assessor = Agent(role="Impact Assessor", goal="Assess the impact of carbon footprint", backstory="Analyzes sustainability scores.", verbose=True)
    sustainability_advisor = Agent(role="Sustainability Advisor", goal="Provide eco-friendly recommendations", backstory="Helps reduce carbon footprint.", verbose=True)

    # CrewAI Tasks
    analyze_task = Task(description="Analyze CO₂ emissions based on user data.", expected_output="Accurate CO₂ footprint estimate.", agent=co2_analyst)
    assess_task = Task(description="Evaluate environmental impact score.", expected_output="Meaningful impact score (0-100).", agent=impact_assessor)
    recommend_task = Task(description="Generate personalized sustainability recommendations.", expected_output="List of actions to reduce footprint.", agent=sustainability_advisor)

    # Forming Crew
    crew = Crew(agents=[co2_analyst, impact_assessor, sustainability_advisor], tasks=[analyze_task, assess_task, recommend_task], process=Process.sequential)
    crew.kickoff(inputs={"total_co2": total_co2})

    # Visualization
    labels = ["Travel", "Energy", "Water", "Diet", "Waste"]
    sizes = [travel_co2, energy_co2, water_co2, diet_co2, waste_co2]
    colors = ["#FF9999", "#66B3FF", "#99FF99", "#FFD700", "#FFCC99"]
    fig, ax = plt.subplots()
    ax.pie(sizes, labels=labels, autopct="%1.1f%%", colors=colors, startangle=140)
    ax.axis("equal")
    st.pyplot(fig)

    # Display Results
    st.markdown(f"### 🌎 Your Estimated Carbon Footprint: **{total_co2} kg CO₂/month**")
    st.markdown(f"### 🔋 Environmental Impact Score: **{score}/100**")
    st.markdown("### 🌱 Recommendations")

    # Generate recommendations based on the score
    # Function to generate recommendations based on the score
    def generate_recommendations(score):
        if score >= 75:
            return {
                "Travel": ["Consider using public transport more often.", "Carpool to reduce emissions."],
                "Energy": ["Switch to renewable energy sources.", "Use energy-efficient appliances."],
                "Diet": ["Reduce meat consumption.", "Buy locally sourced food."],
                "Waste": ["Recycle more.", "Compost organic waste."],
            }
        elif score >= 50:
            return {
                "Travel": ["Limit air travel.", "Use bicycles for short distances."],
                "Energy": ["Turn off lights when not in use.", "Insulate your home to save energy."],
                "Diet": ["Eat more plant-based meals.", "Avoid food waste."],
                "Waste": ["Reduce single-use plastics.", "Donate items instead of discarding."],
            }
        else:
            return {
                "Travel": ["Walk or bike whenever possible.", "Avoid unnecessary trips."],
                "Energy": ["Reduce heating and cooling usage.", "Unplug devices when not in use."],
                "Diet": ["Adopt a vegetarian or vegan diet.", "Grow your own vegetables."],
                "Waste": ["Minimize waste generation.", "Participate in community recycling programs."],
            }
    
    recommendations = generate_recommendations(score)

    # Display recommendations
    for category, tips in recommendations.items():
        st.markdown(f"#### {category}")
        for tip in tips:
            st.markdown(f"- {tip}")

    # Allow user to download a personalized report
    if st.button("Download Report"):
        user_data = {
            "travel": {
                "flights": flights,
                "flight_distance": flight_distance,
                "car_fuel": car_fuel,
                "public_transport_km": public_transport_km,
                "bike_km": bike_km,
                "walking_km": walking_km,
            },
            "energy": {
                "electricity_kwh": electricity_kwh,
                "gas_therms": gas_therms,
                "solar_percentage": solar_percentage,
                "heating_kwh": heating_kwh,
            },
            "water": {"water_liters": water_liters, "hot_water_percentage": hot_water_percentage},
            "diet": {"diet_type": diet_type, "food_waste_kg": food_waste_kg, "plastic_waste_kg": plastic_waste_kg, "local_food_percentage": local_food_percentage},
            "household": {"recycle": recycle, "household_size": household_size, "compost": compost},
            "total_co2": total_co2,
        }

        buffer = create_docx_report(user_data, recommendations, score)
        st.download_button(
            label="📄 Download Your Report",
            data=buffer.getvalue(),
            file_name="carbon_footprint_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )